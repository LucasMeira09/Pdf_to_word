import camelot
from docx import Document
import os

class PdfToWord:
    def __init__(self, path):
        self.path = path
    def get_pdf_text(self):
        # Extraire les tableaux d'un fichier PDF et les exporter en CSV
        tables = camelot.read_pdf(self.path, pages='all')
        csv_path, _ = os.path.splitext(self.path)
        csv_path = csv_path + ".csv"
        tables.export(csv_path, f="csv")

        # rempli la variable avec tout les tableau
        self.dfs = [table.df for table in tables]
        
    def text_to_word(self):
        # recupere les tableau
        dfs = self.dfs
        # la bibliotech
        doc = Document()

        for df in dfs:
            rows, cols = df.shape
            table = doc.add_table(rows=1, cols=cols)
            hdr_cells = table.rows[0].cells

            for n in range(cols):
                hdr_cells[n].text = str(n)

            for l in range(rows):
                row_cells = table.add_row().cells
                for j in range(cols):
                    df_elements = df.iloc[l, j]
                    row_cells[j].text = df_elements
        
            doc.add_paragraph("")

        new_path, _ = os.path.splitext(self.path)
        doc.save(new_path + ".docx")


if __name__ == "__main__":

    path = str(input("The path of your pdf: "))

    pdf_to_word = PdfToWord(path)
    pdf_to_word.get_pdf_text()
    pdf_to_word.text_to_word()